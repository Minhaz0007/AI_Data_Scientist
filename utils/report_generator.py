import pandas as pd
from utils.data_processor import profile_data
import io
from datetime import datetime

def generate_html_report(df, insights=None, title="Data Analysis Report"):
    """
    Generates a comprehensive HTML report for the dataset.
    """
    profile = profile_data(df)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Generate correlation table if available
    corr_html = ""
    if profile.get('correlation'):
        corr_df = pd.DataFrame(profile['correlation'])
        corr_html = f"""
        <h2>Correlation Matrix</h2>
        {corr_df.round(3).to_html(classes='table')}
        """

    # Generate missing values section
    missing_html = ""
    missing_cols = {k: v for k, v in profile['missing_by_col'].items() if v > 0}
    if missing_cols:
        missing_df = pd.DataFrame(list(missing_cols.items()), columns=['Column', 'Missing Count'])
        missing_html = f"""
        <h2>Missing Values</h2>
        {missing_df.to_html(classes='table', index=False)}
        """

    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>{title}</title>
        <style>
            body {{
                font-family: 'Segoe UI', Arial, sans-serif;
                margin: 40px;
                line-height: 1.6;
                color: #333;
            }}
            h1 {{
                color: #2c3e50;
                border-bottom: 3px solid #3498db;
                padding-bottom: 10px;
            }}
            h2 {{
                color: #34495e;
                border-bottom: 2px solid #ecf0f1;
                padding-bottom: 10px;
                margin-top: 30px;
            }}
            .header {{
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 30px;
                border-radius: 10px;
                margin-bottom: 30px;
            }}
            .header h1 {{
                color: white;
                border: none;
                margin: 0;
            }}
            .header p {{
                margin: 5px 0 0 0;
                opacity: 0.9;
            }}
            .metrics {{
                display: flex;
                flex-wrap: wrap;
                gap: 15px;
                margin: 20px 0;
            }}
            .metric {{
                background: #f8f9fa;
                padding: 20px;
                border-radius: 10px;
                min-width: 150px;
                box-shadow: 0 2px 5px rgba(0,0,0,0.1);
            }}
            .metric-value {{
                font-size: 2em;
                font-weight: bold;
                color: #3498db;
            }}
            .metric-label {{
                color: #7f8c8d;
                font-size: 0.9em;
            }}
            .insights-box {{
                background-color: #e8f4f8;
                padding: 25px;
                border-radius: 10px;
                border-left: 4px solid #3498db;
                margin: 20px 0;
                white-space: pre-wrap;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin-top: 20px;
                font-size: 0.9em;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 12px 8px;
                text-align: left;
            }}
            th {{
                background-color: #3498db;
                color: white;
            }}
            tr:nth-child(even) {{
                background-color: #f8f9fa;
            }}
            tr:hover {{
                background-color: #e8f4f8;
            }}
            .footer {{
                margin-top: 50px;
                padding-top: 20px;
                border-top: 1px solid #ecf0f1;
                color: #7f8c8d;
                font-size: 0.9em;
            }}
            @media print {{
                body {{ margin: 20px; }}
                .header {{ break-after: avoid; }}
                table {{ page-break-inside: auto; }}
                tr {{ page-break-inside: avoid; }}
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>{title}</h1>
            <p>Generated on {timestamp}</p>
        </div>

        <h2>Executive Summary</h2>
        <div class="metrics">
            <div class="metric">
                <div class="metric-value">{profile['rows']:,}</div>
                <div class="metric-label">Total Rows</div>
            </div>
            <div class="metric">
                <div class="metric-value">{profile['columns']}</div>
                <div class="metric-label">Columns</div>
            </div>
            <div class="metric">
                <div class="metric-value">{profile['duplicates']:,}</div>
                <div class="metric-label">Duplicate Rows</div>
            </div>
            <div class="metric">
                <div class="metric-value">{profile['missing_total']:,}</div>
                <div class="metric-label">Missing Values</div>
            </div>
        </div>

        <h2>AI-Generated Insights</h2>
        <div class="insights-box">
            {insights if insights else "No AI insights have been generated for this report. Visit the 'AI Insights' section to generate comprehensive analysis."}
        </div>

        <h2>Data Sample (First 10 Rows)</h2>
        {df.head(10).to_html(classes='table', index=False)}

        <h2>Column Information</h2>
        {pd.DataFrame({'Column': df.columns, 'Data Type': df.dtypes.values, 'Non-Null Count': df.count().values, 'Null Count': df.isnull().sum().values}).to_html(classes='table', index=False)}

        {missing_html}

        <h2>Numerical Statistics</h2>
        {pd.DataFrame(profile['numeric_stats']).round(2).to_html(classes='table') if profile['numeric_stats'] else "<p>No numeric columns in the dataset.</p>"}

        {corr_html}

        <div class="footer">
            <p>This report was automatically generated by AI Data Analyst.</p>
            <p>For questions or issues, please contact the data team.</p>
        </div>
    </body>
    </html>
    """
    return html

def generate_pdf_report(df, insights=None, title="Data Analysis Report"):
    """
    Generates a PDF report using reportlab.
    Returns PDF bytes.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    profile = profile_data(df)
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    elements = []
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#2c3e50'),
        spaceAfter=20,
        alignment=TA_CENTER
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#34495e'),
        spaceBefore=20,
        spaceAfter=10
    )

    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=10
    )

    # Title
    elements.append(Paragraph(title, title_style))
    elements.append(Paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    elements.append(Spacer(1, 20))

    # Overview metrics
    elements.append(Paragraph("Dataset Overview", heading_style))

    overview_data = [
        ['Metric', 'Value'],
        ['Total Rows', f'{profile["rows"]:,}'],
        ['Total Columns', str(profile['columns'])],
        ['Duplicate Rows', f'{profile["duplicates"]:,}'],
        ['Missing Values', f'{profile["missing_total"]:,}']
    ]

    overview_table = Table(overview_data, colWidths=[2*inch, 2*inch])
    overview_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9fa')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#ddd'))
    ]))
    elements.append(overview_table)
    elements.append(Spacer(1, 20))

    # AI Insights
    elements.append(Paragraph("AI-Generated Insights", heading_style))
    if insights:
        # Split insights into paragraphs
        for para in insights.split('\n\n'):
            if para.strip():
                elements.append(Paragraph(para.strip(), body_style))
    else:
        elements.append(Paragraph("No AI insights have been generated for this report.", body_style))
    elements.append(Spacer(1, 20))

    # Data sample
    elements.append(Paragraph("Data Sample (First 5 Rows)", heading_style))

    # Prepare data sample table
    sample_df = df.head(5)
    sample_data = [sample_df.columns.tolist()] + sample_df.values.tolist()

    # Truncate long values
    for i, row in enumerate(sample_data):
        sample_data[i] = [str(v)[:30] + '...' if len(str(v)) > 30 else str(v) for v in row]

    # Calculate column widths
    num_cols = len(sample_df.columns)
    col_width = min(1.2*inch, 7*inch/num_cols)

    sample_table = Table(sample_data, colWidths=[col_width]*num_cols)
    sample_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#ddd')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')])
    ]))
    elements.append(sample_table)
    elements.append(Spacer(1, 20))

    # Numerical statistics
    if profile['numeric_stats']:
        elements.append(Paragraph("Numerical Statistics", heading_style))
        stats_df = pd.DataFrame(profile['numeric_stats']).round(2)

        # Limit columns for PDF
        if len(stats_df.columns) > 5:
            stats_df = stats_df.iloc[:, :5]
            elements.append(Paragraph("(Showing first 5 numeric columns)", styles['Italic']))

        stats_data = [['Stat'] + stats_df.columns.tolist()] + [[idx] + row for idx, row in zip(stats_df.index.tolist(), stats_df.values.tolist())]

        stats_table = Table(stats_data)
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#ddd')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')])
        ]))
        elements.append(stats_table)

    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()

def generate_docx_report(df, insights=None, title="Data Analysis Report"):
    """
    Generates a DOCX report using python-docx.
    Returns DOCX bytes.
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    profile = profile_data(df)
    doc = Document()

    # Title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    # Overview
    doc.add_heading('Dataset Overview', level=1)
    overview_table = doc.add_table(rows=5, cols=2)
    overview_table.style = 'Table Grid'

    overview_data = [
        ('Total Rows', f'{profile["rows"]:,}'),
        ('Total Columns', str(profile['columns'])),
        ('Duplicate Rows', f'{profile["duplicates"]:,}'),
        ('Missing Values', f'{profile["missing_total"]:,}'),
        ('Columns', ', '.join(df.columns[:10]) + ('...' if len(df.columns) > 10 else ''))
    ]

    for i, (metric, value) in enumerate(overview_data):
        overview_table.rows[i].cells[0].text = metric
        overview_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # AI Insights
    doc.add_heading('AI-Generated Insights', level=1)
    if insights:
        doc.add_paragraph(insights)
    else:
        doc.add_paragraph("No AI insights have been generated for this report.")

    doc.add_paragraph()

    # Data sample
    doc.add_heading('Data Sample (First 5 Rows)', level=1)
    sample_df = df.head(5)

    # Limit columns for DOCX
    max_cols = 6
    if len(sample_df.columns) > max_cols:
        sample_df = sample_df.iloc[:, :max_cols]
        doc.add_paragraph(f"(Showing first {max_cols} columns)")

    table = doc.add_table(rows=len(sample_df)+1, cols=len(sample_df.columns))
    table.style = 'Table Grid'

    # Headers
    for j, col in enumerate(sample_df.columns):
        table.rows[0].cells[j].text = str(col)

    # Data
    for i, row in enumerate(sample_df.values):
        for j, val in enumerate(row):
            cell_text = str(val)[:50] + '...' if len(str(val)) > 50 else str(val)
            table.rows[i+1].cells[j].text = cell_text

    doc.add_paragraph()

    # Numerical statistics
    if profile['numeric_stats']:
        doc.add_heading('Numerical Statistics', level=1)
        stats_df = pd.DataFrame(profile['numeric_stats']).round(2)

        if len(stats_df.columns) > 5:
            stats_df = stats_df.iloc[:, :5]
            doc.add_paragraph("(Showing first 5 numeric columns)")

        stats_table = doc.add_table(rows=len(stats_df)+1, cols=len(stats_df.columns)+1)
        stats_table.style = 'Table Grid'

        # Headers
        stats_table.rows[0].cells[0].text = 'Statistic'
        for j, col in enumerate(stats_df.columns):
            stats_table.rows[0].cells[j+1].text = str(col)

        # Data
        for i, (idx, row) in enumerate(stats_df.iterrows()):
            stats_table.rows[i+1].cells[0].text = str(idx)
            for j, val in enumerate(row):
                stats_table.rows[i+1].cells[j+1].text = str(round(val, 2))

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
